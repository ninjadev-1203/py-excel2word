from docx import Document
import pandas as pd
from docx.shared import Mm
from xlsxToDocHelper import add_table_to_doc

# read xlsx file
hr_df = pd.read_excel('./test.xlsx')

# select specific columns
hr_df = hr_df[['no', 'name', 'address', 'phonenumber', 'other']]

# create document and set properties
doc = Document()
section = doc.sections[0]
section.left_margin = Mm(5)
section.right_margin = Mm(5)

# add tables
add_table_to_doc(doc, hr_df.iloc[:5],
                 'HR Employee Data 1-5', 'Light Grid Accent 1')
# add second table
add_table_to_doc(
    doc, hr_df.iloc[10:15], 'HR Employee Data 10-15', 'Light Shading Accent 1')

# save to file
doc.save("test.docx")
