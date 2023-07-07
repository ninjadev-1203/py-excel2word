from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def add_table_to_doc(doc, df, heading, table_style='Table Grid'):
    """ Adds a table to a docx document """
    doc.add_heading(
        heading, level=1).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    columns = list(df.columns)
    # add table
    table = doc.add_table(rows=1, cols=len(columns), style=table_style)
    table.autofit = True
    # add columns (if there is '_' then replace with space)
    for col in range(len(columns)):
        set_cell_margins(table.cell(0, col), top=100,
                         start=100, bottom=100, end=50)
        table.cell(0, col).text = columns[col].replace("_", " ").capitalize()
    # add data
    for i, row in enumerate(df.itertuples()):
        table_row = table.add_row().cells
        for col in range(len(columns)):
            set_cell_margins(table_row[col], top=100,
                             start=100, bottom=100, end=50)
            table_row[col].text = str(row[col+1])

    return doc
